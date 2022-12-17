from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture('cat.jpg', width=Inches(2.0))

# name phone number and email details
name = input('What is your name? ')
speak('hello ' + name + 'how are you today?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About Me')

document.add_paragraph(input('Tell me about yourself: '))

# work experience
document.add_heading('Work History')
p = document.add_paragraph()

company = input('Enter Company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experienc at ' + company + ':')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more work history? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experienc at ' + company + ':')
        p.add_run(experience_details)
    else:
        break

        # list of skills
document.add_heading('Skill Set')
skill_set = input('List your primary skill set: ')
p1 = document.add_paragraph(skill_set)
p1.style = 'List Bullet'

# loop until no more skills
while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill_set = input('Enter skill: ')
        p1 = document.add_paragraph(skill_set)
        p1.style = 'List Bullet'
    else:
        break

        # footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "this is a footer that could be hard coded into document"


document.save('cv.docx')
