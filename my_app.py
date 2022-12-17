from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture('cat.jpg', width=Inches(2.0))

# name phone number and email details
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About Me')

document.add_paragraph(input('Tell me about yourself: '))

# work experience
document.add_heading('Work History')
p = document.add_paragraph()

company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experienc at ' + company)
p.add_run(experience_details)

document.save('cv.docx')
